from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

for section in doc.sections:
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2)
    section.right_margin = Cm(2)

def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def add_heading(doc, text, level=1):
    return doc.add_heading(text, level=level)

def add_table_with_header(doc, headers, rows, col_widths=None, header_color='2E4057'):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        hdr_cells[i].paragraphs[0].runs[0].bold = True
        hdr_cells[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        hdr_cells[i].paragraphs[0].runs[0].font.size = Pt(9)
        hdr_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_bg(hdr_cells[i], header_color)
    for ri, row_data in enumerate(rows):
        cells = table.rows[ri + 1].cells
        for ci, val in enumerate(row_data):
            cells[ci].text = str(val)
            cells[ci].paragraphs[0].runs[0].font.size = Pt(9)
            if ri % 2 == 1:
                set_cell_bg(cells[ci], 'F5F7FA')
    if col_widths:
        for row in table.rows:
            for i, w in enumerate(col_widths):
                row.cells[i].width = Cm(w)
    return table

# ===== Title =====
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('Entry Sheet Detail Screen — Field Specification')
run.bold = True
run.font.size = Pt(18)
run.font.color.rgb = RGBColor(0x2E, 0x40, 0x57)

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub.add_run('Japanese proper nouns (ステータス, ラベル, etc.) are kept as-is.').font.size = Pt(10)

doc.add_paragraph()

# ===== Overview =====
add_heading(doc, 'Overview', level=2)
p = doc.add_paragraph(
    'The Entry Sheet detail screen allows manufacturers (メーカー) to submit product entry requests to a retailer (小売店). '
    'Each sheet follows a workflow: 下書き (draft) → 完了 (completed) → review → 承認済み (approved).\n'
    'The screen is divided into 6 sections: 作成情報, 詳細情報, 変更履歴, 商品情報, 販促物情報, and 管理者メモ.'
)
p.runs[0].font.size = Pt(10)
doc.add_paragraph()

# ===== ステータス =====
add_heading(doc, 'ステータス (Status) Definitions', level=2)
add_table_with_header(doc,
    ['ステータス value', 'Display name', 'Description'],
    [
        ['draft',               '下書き',       'Saved as a draft, not yet submitted'],
        ['completed',           '完了',         'Submitted with product images'],
        ['completed_no_image',  '画像なし完了', 'Submitted without product images'],
        ['revision_requested',  '修正依頼',     'Admin has requested corrections'],
        ['approved',            '承認済み',     'Approved by admin'],
    ],
    col_widths=[4, 3.5, 12.5]
)
doc.add_paragraph()

# ===== User Roles =====
add_heading(doc, 'User Roles', level=2)
add_table_with_header(doc,
    ['Role', 'Description'],
    [
        ['ADMIN',    'Administrator — can approve, request revisions, and edit all fields including 管理者メモ'],
        ['STAFF',    'Manufacturer staff — creates and submits sheets'],
        ['RETAILER', 'Retailer — read-only access'],
    ],
    col_widths=[3.5, 16.5]
)
doc.add_paragraph()

doc.add_page_break()

# ===== 1. Field Spec Table =====
add_heading(doc, '1. Entry Sheet Detail Screen — Field Specification Table', level=1)

p = doc.add_paragraph('★ Mandatory fields: タイトル, 展開期間, 棚割り幅, 変更履歴, リスク分類')
p.runs[0].bold = True
p.runs[0].font.color.rgb = RGBColor(0xC0, 0x39, 0x2B)
doc.add_paragraph()

field_headers = ['Section', 'Field label (UI)', 'Field key', 'Data type', 'Input method', 'Required\n(下書き)', 'Required\n(完了時)', 'Validation rule', 'Error message']
field_col_widths = [2.8, 3.5, 3.8, 2.2, 3.2, 1.3, 1.3, 5.2, 6.7]

field_rows = [
    # 作成情報
    ['作成情報', '作成日時\n(Created at)',      'createdAt',    'string (ISO date)', 'Auto-filled, read-only', '−', '−', '−', '−'],
    ['作成情報', '更新日時\n(Updated at)',      'updatedAt',    'string (ISO date)', 'Auto-filled, read-only', '−', '−', '−', '−'],
    ['作成情報', 'メーカー名\n(Manufacturer)', 'manufacturerName', 'string',         'Auto-filled, read-only', '−', '−', '−', '−'],
    ['作成情報', '作成者\n(Creator name)',      'creatorName',  'string',            'Text input',             '○', '○', 'Required', '作成者を入力してください'],
    ['作成情報', '作成者メール\n(Email)',        'email',        'string',            'Text input',             '○', '○', 'Required', '作成者メールを入力してください'],
    ['作成情報', '作成者電話番号\n(Phone)',      'phoneNumber',  'string',            'Text input',             '○', '○', 'Required', '作成者電話番号を入力してください'],
    # 詳細情報
    ['詳細情報 ★', '★ タイトル\n(Title)',        'title',        'string',            'Text input',             '○', '○', 'Required', 'タイトルを入力してください'],
    ['詳細情報',   'カテゴリ名\n(Category)',     'shelfName',    'string',            'Dropdown (master list)', '−', '−', 'Must match master', '−'],
    ['詳細情報',   '提出先\n(Submission dest.)', 'caseName',     'string',            'Dropdown (master list)', '−', '−', 'Must match master', '−'],
    ['詳細情報',   '案件\n(Project)',            'project',      'string',            'Dropdown (master list)', '−', '−', 'Must match master', '−'],
    ['詳細情報 ★', '★ 展開期間 スタート月\n(Deployment start month)', 'deploymentStartMonth', 'number (1–12)', 'Dropdown (current + 3 months)', '−', '−', 'Integer 1–12', '−'],
    ['詳細情報 ★', '★ 展開期間 終了月\n(Deployment end month)',       'deploymentEndMonth',   'number (1–12)', 'Auto-calculated; ADMIN can override', '−', '−', 'Auto-set to startMonth + 1', '−'],
    ['詳細情報 ★', '★ 棚割り幅 ラベル\n(Face option label)',          'faceLabel',    'string',   'Dropdown (per-manufacturer master)', '−', '−', 'Must match manufacturer フェイスオプション', '−'],
    ['詳細情報 ★', '★ 棚割り幅 MAX幅\n(Face max width)',              'faceMaxWidth', 'number (mm)', 'Auto-filled from selected ラベル', '−', '−', 'Σ(width × facingCount) ≤ faceMaxWidth', '商品幅合計がフェイスMAX値（{MAX}mm）を超えているため完了できません。'],
    ['詳細情報',   'エントリシート補足情報\n(Sheet notes)',           'notes',        'string',   'Textarea',                          '−', '−', '−', '−'],
    ['詳細情報',   '添付ファイル\n(Attachments)',                     'attachments',  'Attachment[]', 'File upload',                   '−', '−', 'Max 25 MB per file', 'ファイルサイズは25MB以下にしてください: {filename}'],
    # 商品情報
    ['商品情報',   'JANコード\n(JAN code)',         'janCode',             'string',       'Text input',                   '○', '○', '8, 13, or 16 digits only', '(JAN format error)'],
    ['商品情報',   '商品名\n(Product name)',         'productName',         'string',       'Text input',                   '○', '○', 'Required', '−'],
    ['商品情報',   '商品画像\n(Product image)',      'productImage',        'string\n(URL/Base64)', 'Image upload',         '−', '○', 'Max 25 MB; min 1000px short side; AI/PNG/JPEG/EPS', '−'],
    ['商品情報 ★', '★ リスク分類\n(Risk classification)', 'riskClassification', 'string',  'Dropdown (master list)',        '○', '○', 'Required; must match master', '−'],
    ['商品情報',   '特定成分\n(Specific ingredients)', 'specificIngredients', 'string[]', 'Checkboxes (multi-select)',     '−', '−', 'Must match master', '−'],
    ['商品情報',   '幅 (Width, mm)',                 'width',               'number',       'Number input',                 '○', '○', 'Integer > 0', '−'],
    ['商品情報',   '高さ (Height, mm)',               'height',              'number',       'Number input',                 '○', '○', 'Integer > 0', '−'],
    ['商品情報',   '奥行 (Depth, mm)',                'depth',               'number',       'Number input',                 '○', '○', 'Integer > 0', '−'],
    ['商品情報',   'フェイシング数\n(Facing count)',  'facingCount',         'number',       'Number input',                 '○', '○', 'Integer > 0', '−'],
    ['商品情報',   '納品日\n(Arrival date)',           'arrivalDate',         'string\n(YYYY-MM-DD)', 'Date picker',          '−', '−', 'YYYY-MM-DD format', '−'],
    ['商品情報',   'キャッチコピー\n(Catch copy)',    'catchCopy',           'string',       'Text input',                   '−', '−', '−', '−'],
    ['商品情報',   '補足事項\n(Product notes)',       'productNotes',        'string',       'Textarea',                     '−', '−', '−', '−'],
    ['商品情報',   '添付ファイル（商品）\n(Product attachments)', 'productAttachments', 'Attachment[]', 'File upload',       '−', '−', 'Max 25 MB per file', 'ファイルサイズは25MB以下にしてください: {filename}'],
    # 販促物情報
    ['販促物情報', '販促物有無\n(Has promo material)', 'hasPromoMaterial', "'yes'|'no'", 'Radio button', '−', '−', '−', '−'],
    ['販促物情報', '香り見本・陳列売什器\n(Special fixture)', 'specialFixture', 'string', 'Text input', '−', '−', '−', '−'],
    ['販促物情報', '幅 (Promo width, mm)',   'promoWidth',  'number', 'Number input', '−', '−', '−', '−'],
    ['販促物情報', '高さ (Promo height, mm)', 'promoHeight', 'number', 'Number input', '−', '−', '−', '−'],
    ['販促物情報', '奥行 (Promo depth, mm)', 'promoDepth',  'number', 'Number input', '−', '−', '−', '−'],
    ['販促物情報', '販促物画像\n(Promo image)', 'promoImage', 'string (URL)', 'Image upload', '−', '−', 'Max 25 MB', '−'],
    ['販促物情報', '納品日\n(Delivery date)', 'deliveryDate', 'string\n(YYYY-MM-DD)', 'Date picker', '−', '−', 'YYYY-MM-DD format', '−'],
    # 管理者メモ
    ['管理者メモ\n(ADMIN only)', '販促CD\n(Promo code)',                    'promoCode',        'string', 'Text input', '−', '−', '−', '−'],
    ['管理者メモ\n(ADMIN only)', 'ボードピッキングJAN\n(Board picking JAN)', 'boardPickingJan',  'string', 'Text input', '−', '−', '−', '−'],
    ['管理者メモ\n(ADMIN only)', '期限表URL\n(Deadline table URL)',          'deadlineTableUrl', 'string', 'Text input', '−', '−', 'Max length enforced', '期限表URLは{N}文字以内で入力してください'],
    ['管理者メモ\n(ADMIN only)', '帯パターン\n(Band pattern)',              'bandPattern',      'string', 'Text input', '−', '−', 'Max length enforced', '帯パターンは{N}文字以内で入力してください'],
    ['管理者メモ\n(ADMIN only)', '対象店舗数\n(Target store count)',        'targetStoreCount', 'number', 'Number input', '−', '−', '−', '−'],
    ['管理者メモ\n(ADMIN only)', '印刷依頼数量 ボード①',                   'printBoard1Count', 'number', 'Number input', '−', '−', '−', '−'],
    ['管理者メモ\n(ADMIN only)', '印刷依頼数量 ボード②',                   'printBoard2Count', 'number', 'Number input', '−', '−', '−', '−'],
    ['管理者メモ\n(ADMIN only)', '印刷依頼数量 帯①',                       'printBand1Count',  'number', 'Number input', '−', '−', '−', '−'],
    ['管理者メモ\n(ADMIN only)', '印刷依頼数量 帯②',                       'printBand2Count',  'number', 'Number input', '−', '−', '−', '−'],
    ['管理者メモ\n(ADMIN only)', '印刷依頼数量 その他\n(Print qty - other)', 'printOther',      'string', 'Text input', '−', '−', 'Max length enforced', '印刷依頼数量 その他は{N}文字以内で入力してください'],
    ['管理者メモ\n(ADMIN only)', '備品\n(Equipment note)',                  'equipmentNote',    'string', 'Textarea', '−', '−', 'Max length enforced', '備品は{N}文字以内で入力してください'],
    ['管理者メモ\n(ADMIN only)', '備考\n(Admin note)',                      'adminNote',        'string', 'Textarea', '−', '−', 'Max length enforced', '備考は{N}文字以内で入力してください'],
]
add_table_with_header(doc, field_headers, field_rows, col_widths=field_col_widths)

doc.add_paragraph()
p = doc.add_paragraph('Note — 商品情報 tab status badge:')
p.runs[0].bold = True
p.runs[0].font.size = Pt(10)
add_table_with_header(doc,
    ['Badge', 'Condition'],
    [
        ['完了',   'All required fields filled'],
        ['入力中', 'Some fields filled, some missing'],
        ['未入力', 'No fields filled yet'],
    ],
    col_widths=[3, 17]
)

doc.add_paragraph()
p = doc.add_paragraph('棚割り幅 constraint (shelf width check):')
p.runs[0].bold = True
p.runs[0].font.size = Pt(10)
p = doc.add_paragraph('  shelfWidthTotal = Σ (product.width × product.facingCount)')
p.runs[0].font.name = 'Courier New'
p.runs[0].font.size = Pt(9)
p = doc.add_paragraph('  If faceLabel is selected → shelfWidthTotal must not exceed faceMaxWidth.')
p.runs[0].font.size = Pt(9)

doc.add_page_break()

# ===== 2. 変更履歴 Display Rules =====
add_heading(doc, '2. 変更履歴 (Change History) — Display Rules & Logic', level=1)

p = doc.add_paragraph(
    '変更履歴 entries are auto-recorded on every save and review action. '
    'Up to 30 entries are retained per sheet (keepLatestCount: 30); oldest entries are deleted automatically.'
)
p.runs[0].font.size = Pt(10)
doc.add_paragraph()

# 2-1
add_heading(doc, '2-1. When is a history entry recorded?', level=2)

add_heading(doc, 'Trigger A — On sheet save  (PUT /api/sheets/[id])', level=3)
p = doc.add_paragraph(
    'The function buildRevisionSummary(before, after) compares the sheet data before and after saving '
    'and generates the summary string automatically.'
)
p.runs[0].font.size = Pt(10)

add_table_with_header(doc,
    ['Case', 'Generated summary text', 'Icon pattern'],
    [
        ['First-time creation\n(no previous data)',
         '新規作成: タイトル="{title}" / 商品件数={n}',
         '⊕ Sky blue (新規作成)'],
        ['Changes detected\n(diff exists)',
         'Each changed field on a new line:\n{label}: {before} -> {after}\n(up to 80 items)',
         'Determined by content — see icon rules below'],
        ['No changes\n(data identical)',
         '変更なしで保存',
         '📄 Gray (default)'],
    ],
    col_widths=[4.5, 9.5, 6]
)
doc.add_paragraph()

p = doc.add_paragraph('Fields compared by buildRevisionSummary():')
p.runs[0].bold = True
p.runs[0].font.size = Pt(10)
add_table_with_header(doc,
    ['Log label in summary', 'Field key', 'Notes'],
    [
        ['タイトル',         'title',                        ''],
        ['案件',             'caseName',                     ''],
        ['補足',             'notes',                        ''],
        ['カテゴリ名',       'shelfName',                    ''],
        ['作成者名',         'creatorName',                  ''],
        ['作成者メール',     'email',                        ''],
        ['作成者電話',       'phoneNumber',                  ''],
        ['状態',             'status',                       'e.g. draft -> completed'],
        ['展開スタート月',   'deploymentStartMonth',         ''],
        ['展開終了月',       'deploymentEndMonth',           ''],
        ['棚割り幅',         'faceLabel',                    ''],
        ['フェイスMAX値',    'faceMaxWidth',                 ''],
        ['商品件数',         'products.length',              'Only recorded when count changes'],
        ['商品N.商品名',     'products[n].productName',      'N = product index (1-based)'],
        ['商品N.JAN',        'products[n].janCode',          ''],
        ['商品N.リスク分類', 'products[n].riskClassification', ''],
    ],
    col_widths=[4.5, 5.5, 10]
)
doc.add_paragraph()

add_heading(doc, 'Trigger B — On admin review action  (POST /api/sheets/[id]/review)', level=3)
add_table_with_header(doc,
    ['Admin action', 'Generated summary text', 'Icon pattern'],
    [
        ['Approve (承認)',                  '承認しました',               '✔ Emerald (承認・完了)'],
        ['Request revision (修正依頼)',     '修正依頼: {comment text}',   '⚠ Orange (修正依頼)'],
    ],
    col_widths=[5.5, 7.5, 7]
)
doc.add_paragraph()

# 2-2
add_heading(doc, '2-2. Icon assignment rules  (getRevisionIcon function)', level=2)
p = doc.add_paragraph(
    'Icons are determined by keyword matching against the summary string. '
    'Rules are evaluated top-to-bottom; the first match wins.'
)
p.runs[0].font.size = Pt(10)

add_table_with_header(doc,
    ['Priority', 'Keyword(s) in summary', 'Icon (lucide-react)', 'Color (Tailwind)'],
    [
        ['1', '修正依頼',              '⚠  AlertTriangle', 'text-orange-600  bg-orange-100'],
        ['2', 'ステータス / 状態 / →', '↺  RefreshCw',    'text-amber-500   bg-amber-50'],
        ['3', '承認 / 確定 / 完了',    '✔  CheckCircle',  'text-emerald-500 bg-emerald-50'],
        ['4', '差戻 / 却下 / 返却',    '↩  RotateCcw',    'text-rose-500    bg-rose-50'],
        ['5', '商品 / product',        '📦 Package',      'text-violet-500  bg-violet-50'],
        ['6', '作成 / 新規 / 追加',    '⊕  PlusCircle',   'text-sky-500     bg-sky-50'],
        ['7', '編集 / 更新 / 変更',    '✎  Edit3',        'text-blue-500    bg-blue-50'],
        ['8', '(no match)',            '📄 FileText',     'text-slate-400   bg-slate-100'],
    ],
    col_widths=[1.8, 6, 4.2, 8]
)
doc.add_paragraph()

# 2-3
add_heading(doc, '2-3. Screen layout & display spec', level=2)
for item in [
    'Each entry layout:  [ icon ]  Date & time (ja-JP locale)   Changed-by name  /  summary text (multi-line supported)',
    'Entry count shown as a badge in the section heading — e.g. "変更履歴（直近） [3件]"',
    'List max height: 288 px — scrollable when overflow',
    'When no entries exist: displays "履歴はまだありません。"',
    'Retention limit: 30 entries per sheet (oldest auto-deleted)',
]:
    p = doc.add_paragraph(item, style='List Bullet')
    p.runs[0].font.size = Pt(10)

doc.add_paragraph()
add_heading(doc, 'Display example', level=3)
example = (
    "変更履歴（直近）  [3件]\n"
    "─────────────────────────────────────────────\n"
    "⊕  2025/04/01 10:00  田中 太郎\n"
    "   新規作成: タイトル=\"春の新商品\" / 商品件数=1\n"
    "─────────────────────────────────────────────\n"
    "✎  2025/04/02 14:30  田中 太郎\n"
    "   タイトル: 春の新商品 -> 春の新商品キャンペーン\n"
    "   展開スタート月: 4 -> 5\n"
    "─────────────────────────────────────────────\n"
    "⚠  2025/04/03 09:15  管理者\n"
    "   修正依頼: 商品画像を添付してください"
)
p = doc.add_paragraph()
run = p.add_run(example)
run.font.name = 'Courier New'
run.font.size = Pt(9)

doc.add_paragraph()

# ===== API reference =====
add_heading(doc, '3. API Endpoints (reference)', level=1)
add_table_with_header(doc,
    ['Method', 'Path', 'Description'],
    [
        ['GET',  '/api/sheets/[id]',          'Fetch sheet detail'],
        ['PUT',  '/api/sheets/[id]',          'Save / update sheet (triggers buildRevisionSummary)'],
        ['POST', '/api/sheets/[id]/review',   'Admin approve or request_revision (fixed summary text)'],
        ['GET',  '/api/sheets/[id]/revisions','Fetch 変更履歴 list'],
    ],
    col_widths=[2, 6, 12]
)

out_path = '/home/user/medicom-entry-sheet/docs/Entry_Sheet_Detail_Screen_Spec.docx'
doc.save(out_path)
print(f'Saved: {out_path}')
