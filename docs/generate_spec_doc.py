from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

# Page margins
for section in doc.sections:
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2)
    section.right_margin = Cm(2)

def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def add_heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    return p

def add_table_with_header(doc, headers, rows, col_widths=None, header_color='2E4057'):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    # Header row
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        hdr_cells[i].paragraphs[0].runs[0].bold = True
        hdr_cells[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        hdr_cells[i].paragraphs[0].runs[0].font.size = Pt(9)
        hdr_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_bg(hdr_cells[i], header_color)
    # Data rows
    for ri, row_data in enumerate(rows):
        cells = table.rows[ri + 1].cells
        for ci, val in enumerate(row_data):
            cells[ci].text = str(val)
            cells[ci].paragraphs[0].runs[0].font.size = Pt(9)
            if ri % 2 == 1:
                set_cell_bg(cells[ci], 'F5F7FA')
    if col_widths:
        for row in table.rows:
            for i, w in enumerate(col_widths):
                row.cells[i].width = Cm(w)
    return table

# ===== Title =====
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('エントリーシート詳細画面 仕様書')
run.bold = True
run.font.size = Pt(18)
run.font.color.rgb = RGBColor(0x2E, 0x40, 0x57)

doc.add_paragraph()

# ===== ステータス定義 =====
add_heading(doc, 'ステータス定義', level=2)

status_headers = ['ステータス値', '表示名', '説明']
status_rows = [
    ['draft',               '下書き',       '保存済み・未提出'],
    ['completed',           '完了',         '商品画像あり・提出済み'],
    ['completed_no_image',  '画像なし完了', '商品画像なし・提出済み'],
    ['revision_requested',  '修正依頼',     '管理者が修正を要求した状態'],
    ['approved',            '承認済み',     '管理者が承認した状態'],
]
add_table_with_header(doc, status_headers, status_rows, col_widths=[4, 3.5, 9])
doc.add_paragraph()

# ===== 1. フィールド仕様表 =====
add_heading(doc, '1. エントリーシート詳細画面 フィールド仕様表', level=1)

p = doc.add_paragraph('★ マスト5項目：タイトル、展開期間、棚割り幅、変更履歴、リスク分類')
p.runs[0].bold = True
p.runs[0].font.color.rgb = RGBColor(0xC0, 0x39, 0x2B)
doc.add_paragraph()

field_headers = ['セクション', 'フィールド名（UI）', 'フィールドキー', 'データ型', '入力方式', '必須\n(下書き)', '必須\n(完了時)', 'バリデーションルール', 'エラーメッセージ']
field_col_widths = [2.8, 3.5, 3.8, 2.2, 3.2, 1.2, 1.2, 5.5, 5.5]

field_rows = [
    # 作成情報
    ['作成情報', '作成日時',     'createdAt',    'string (ISO date)', '自動入力（読み取り専用）', '−', '−', '−', '−'],
    ['作成情報', '更新日時',     'updatedAt',    'string (ISO date)', '自動入力（読み取り専用）', '−', '−', '−', '−'],
    ['作成情報', 'メーカー名',   'manufacturerName', 'string',        '自動入力（読み取り専用）', '−', '−', '−', '−'],
    ['作成情報', '作成者',       'creatorName',  'string',            'テキスト入力',             '○', '○', '必須',  '作成者を入力してください'],
    ['作成情報', '作成者メール', 'email',         'string',           'テキスト入力',             '○', '○', '必須',  '作成者メールを入力してください'],
    ['作成情報', '作成者電話番号','phoneNumber',  'string',           'テキスト入力',             '○', '○', '必須',  '作成者電話番号を入力してください'],
    # 詳細情報
    ['詳細情報 ★', '★ タイトル', 'title',         'string',          'テキスト入力',             '○', '○', '必須',  'タイトルを入力してください'],
    ['詳細情報',  'カテゴリ名',  'shelfName',    'string',            'プルダウン（マスター）',   '−', '−', 'マスターから選択', '−'],
    ['詳細情報',  '提出先',      'caseName',     'string',            'プルダウン（マスター）',   '−', '−', 'マスターから選択', '−'],
    ['詳細情報',  '案件',        'project',      'string',            'プルダウン（マスター）',   '−', '−', 'マスターから選択', '−'],
    ['詳細情報 ★', '★ 展開期間（スタート月）', 'deploymentStartMonth', 'number (1〜12)', 'プルダウン（当月〜3ヶ月先）', '−', '−', '1〜12の整数', '−'],
    ['詳細情報 ★', '★ 展開期間（終了月）',    'deploymentEndMonth',   'number (1〜12)', '自動計算 / 管理者のみ上書き可', '−', '−', 'スタート月+1を自動設定', '−'],
    ['詳細情報 ★', '★ 棚割り幅（ラベル）',    'faceLabel',           'string',         'プルダウン（メーカー別マスター）', '−', '−', 'メーカー別フェイスオプションから選択', '−'],
    ['詳細情報 ★', '★ 棚割り幅（MAX幅）',     'faceMaxWidth',        'number (mm)',    '自動入力（ラベル選択に連動）', '−', '−', '商品幅合計 ≦ MAX幅', '商品幅合計がフェイスMAX値（{MAX}mm）を超えているため完了できません。'],
    ['詳細情報',  'エントリシート補足情報', 'notes', 'string',        'テキストエリア',           '−', '−', '−', '−'],
    ['詳細情報',  '添付ファイル', 'attachments', 'Attachment[]',      'ファイルアップロード',     '−', '−', '1ファイルあたり25MB以下', 'ファイルサイズは25MB以下にしてください: {ファイル名}'],
    # 商品情報
    ['商品情報',  'JANコード',   'janCode',      'string',            'テキスト入力',             '○', '○', '8・13・16桁の数字のみ', '（JANコード形式エラー）'],
    ['商品情報',  '商品名',      'productName',  'string',            'テキスト入力',             '○', '○', '必須', '−'],
    ['商品情報',  '商品画像',    'productImage', 'string (URL/Base64)','画像アップロード',         '−', '○', '25MB以下 / 短辺1000px以上 / AI・PNG・JPEG・EPS', '−'],
    ['商品情報 ★', '★ リスク分類', 'riskClassification', 'string',  'プルダウン（マスター）',   '○', '○', '必須・マスターから選択', '−'],
    ['商品情報',  '特定成分',    'specificIngredients', 'string[]',  'チェックボックス（複数）', '−', '−', 'マスターから選択', '−'],
    ['商品情報',  '幅（mm）',    'width',        'number',            '数値入力',                 '○', '○', '0より大きい整数', '−'],
    ['商品情報',  '高さ（mm）',  'height',       'number',            '数値入力',                 '○', '○', '0より大きい整数', '−'],
    ['商品情報',  '奥行（mm）',  'depth',        'number',            '数値入力',                 '○', '○', '0より大きい整数', '−'],
    ['商品情報',  'フェイシング数', 'facingCount','number',           '数値入力',                 '○', '○', '0より大きい整数', '−'],
    ['商品情報',  '納品日',      'arrivalDate',  'string (YYYY-MM-DD)','日付ピッカー',            '−', '−', 'YYYY-MM-DD形式', '−'],
    ['商品情報',  'キャッチコピー','catchCopy',  'string',            'テキスト入力',             '−', '−', '−', '−'],
    ['商品情報',  '補足事項',    'productNotes', 'string',            'テキストエリア',           '−', '−', '−', '−'],
    ['商品情報',  '添付ファイル（商品）', 'productAttachments', 'Attachment[]', 'ファイルアップロード', '−', '−', '1ファイルあたり25MB以下', 'ファイルサイズは25MB以下にしてください: {ファイル名}'],
    # 販促物情報
    ['販促物情報', '販促物有無',  'hasPromoMaterial', "'yes'|'no'",  'ラジオボタン',             '−', '−', '−', '−'],
    ['販促物情報', '香り見本・陳列売什器', 'specialFixture', 'string','テキスト入力',           '−', '−', '−', '−'],
    ['販促物情報', '幅（mm）',    'promoWidth',   'number',           '数値入力',                 '−', '−', '−', '−'],
    ['販促物情報', '高さ（mm）',  'promoHeight',  'number',           '数値入力',                 '−', '−', '−', '−'],
    ['販促物情報', '奥行（mm）',  'promoDepth',   'number',           '数値入力',                 '−', '−', '−', '−'],
    ['販促物情報', '販促物画像',  'promoImage',   'string (URL)',     '画像アップロード',         '−', '−', '25MB以下', '−'],
    ['販促物情報', '納品日',      'deliveryDate', 'string (YYYY-MM-DD)','日付ピッカー',          '−', '−', 'YYYY-MM-DD形式', '−'],
    # 管理者メモ
    ['管理者メモ\n(ADMIN only)', '販促CD', 'promoCode', 'string', 'テキスト入力', '−', '−', '−', '−'],
    ['管理者メモ\n(ADMIN only)', 'ボードピッキングJAN', 'boardPickingJan', 'string', 'テキスト入力', '−', '−', '−', '−'],
    ['管理者メモ\n(ADMIN only)', '期限表URL', 'deadlineTableUrl', 'string', 'テキスト入力', '−', '−', '文字数上限あり', '期限表URLは{N}文字以内で入力してください'],
    ['管理者メモ\n(ADMIN only)', '帯パターン', 'bandPattern', 'string', 'テキスト入力', '−', '−', '文字数上限あり', '帯パターンは{N}文字以内で入力してください'],
    ['管理者メモ\n(ADMIN only)', '対象店舗数', 'targetStoreCount', 'number', '数値入力', '−', '−', '−', '−'],
    ['管理者メモ\n(ADMIN only)', '印刷依頼数量 ボード①', 'printBoard1Count', 'number', '数値入力', '−', '−', '−', '−'],
    ['管理者メモ\n(ADMIN only)', '印刷依頼数量 ボード②', 'printBoard2Count', 'number', '数値入力', '−', '−', '−', '−'],
    ['管理者メモ\n(ADMIN only)', '印刷依頼数量 帯①', 'printBand1Count', 'number', '数値入力', '−', '−', '−', '−'],
    ['管理者メモ\n(ADMIN only)', '印刷依頼数量 帯②', 'printBand2Count', 'number', '数値入力', '−', '−', '−', '−'],
    ['管理者メモ\n(ADMIN only)', '印刷依頼数量 その他', 'printOther', 'string', 'テキスト入力', '−', '−', '文字数上限あり', '印刷依頼数量 その他は{N}文字以内で入力してください'],
    ['管理者メモ\n(ADMIN only)', '備品', 'equipmentNote', 'string', 'テキストエリア', '−', '−', '文字数上限あり', '備品は{N}文字以内で入力してください'],
    ['管理者メモ\n(ADMIN only)', '備考', 'adminNote', 'string', 'テキストエリア', '−', '−', '文字数上限あり', '備考は{N}文字以内で入力してください'],
]
add_table_with_header(doc, field_headers, field_rows, col_widths=field_col_widths)

doc.add_page_break()

# ===== 2. 変更履歴の表示ルール =====
add_heading(doc, '2. 変更履歴の表示ルール', level=1)

p = doc.add_paragraph('変更履歴（EntrySheetRevision）は保存・レビューのたびに自動記録されます。最大30件保持し、古いものから削除されます。')
p.runs[0].font.size = Pt(10)
doc.add_paragraph()

# 2-1 記録タイミング
add_heading(doc, '2-1. 記録タイミングとsummary生成ロジック', level=2)

add_heading(doc, 'トリガーA：シート保存時（PUT /api/sheets/[id]）', level=3)

p = doc.add_paragraph('保存前後のシートデータを比較し、buildRevisionSummary() で自動生成します。')
p.runs[0].font.size = Pt(10)

trigger_headers = ['ケース', 'summaryの内容（記録テキスト）', '対応アイコンパターン']
trigger_rows = [
    ['新規作成（保存前データなし）',
     '新規作成: タイトル="{title}" / 商品件数={n}',
     '⊕ 新規作成（スカイ色）'],
    ['変更あり（差分あり）',
     '変わった項目を改行連結:\n{ラベル}: {変更前} -> {変更後}\n（最大80項目）',
     '項目内容で自動判定（下表参照）'],
    ['変更なし（内容が同一）',
     '変更なしで保存',
     '📄 デフォルト（グレー）'],
]
add_table_with_header(doc, trigger_headers, trigger_rows, col_widths=[5.5, 9, 5.5])
doc.add_paragraph()

p = doc.add_paragraph('比較されるフィールド一覧：')
p.runs[0].bold = True
p.runs[0].font.size = Pt(10)

compare_headers = ['summaryに出力されるラベル', '対象フィールドキー', '備考']
compare_rows = [
    ['タイトル',         'title',                   ''],
    ['案件',             'caseName',                ''],
    ['補足',             'notes',                   ''],
    ['カテゴリ名',       'shelfName',               ''],
    ['作成者名',         'creatorName',             ''],
    ['作成者メール',     'email',                   ''],
    ['作成者電話',       'phoneNumber',             ''],
    ['状態',             'status',                  '例: draft -> completed'],
    ['展開スタート月',   'deploymentStartMonth',    ''],
    ['展開終了月',       'deploymentEndMonth',      ''],
    ['棚割り幅',         'faceLabel',               ''],
    ['フェイスMAX値',    'faceMaxWidth',            ''],
    ['商品件数',         'products.length',         '増減があった場合のみ記録'],
    ['商品N.商品名',     'products[n].productName', 'Nは商品インデックス（1始まり）'],
    ['商品N.JAN',        'products[n].janCode',     ''],
    ['商品N.リスク分類', 'products[n].riskClassification', ''],
]
add_table_with_header(doc, compare_headers, compare_rows, col_widths=[4.5, 5.5, 9])
doc.add_paragraph()

add_heading(doc, 'トリガーB：レビューアクション時（POST /api/sheets/[id]/review）', level=3)

review_headers = ['管理者の操作', 'summaryに記録される内容', '対応アイコンパターン']
review_rows = [
    ['承認（approve）',         '承認しました',          '✔ 承認・完了（エメラルド）'],
    ['修正依頼（request_revision）', '修正依頼: {コメント本文}', '⚠ 修正依頼（オレンジ）'],
]
add_table_with_header(doc, review_headers, review_rows, col_widths=[5.5, 7, 7.5])
doc.add_paragraph()

# 2-2 アイコン判定ルール
add_heading(doc, '2-2. アイコン判定ルール（getRevisionIcon関数）', level=2)

p = doc.add_paragraph('summaryテキストのキーワードマッチで判定。上から優先順に評価し、最初にマッチしたものを採用。')
p.runs[0].font.size = Pt(10)

icon_headers = ['優先順', '判定キーワード（summaryに含む）', 'アイコン', '色（Tailwind）']
icon_rows = [
    ['1', '修正依頼',               '⚠ AlertTriangle', 'text-orange-600 bg-orange-100'],
    ['2', 'ステータス / 状態 / →',  '↺ RefreshCw',    'text-amber-500 bg-amber-50'],
    ['3', '承認 / 確定 / 完了',     '✔ CheckCircle',  'text-emerald-500 bg-emerald-50'],
    ['4', '差戻 / 却下 / 返却',     '↩ RotateCcw',    'text-rose-500 bg-rose-50'],
    ['5', '商品 / product',         '📦 Package',     'text-violet-500 bg-violet-50'],
    ['6', '作成 / 新規 / 追加',     '⊕ PlusCircle',   'text-sky-500 bg-sky-50'],
    ['7', '編集 / 更新 / 変更',     '✎ Edit3',        'text-blue-500 bg-blue-50'],
    ['8', '（どれにも該当しない）',  '📄 FileText',    'text-slate-400 bg-slate-100'],
]
add_table_with_header(doc, icon_headers, icon_rows, col_widths=[1.8, 6, 4, 8.2])
doc.add_paragraph()

# 2-3 表示レイアウト
add_heading(doc, '2-3. 画面表示レイアウト', level=2)

layout_items = [
    '各エントリの表示構造：[ アイコン ]  日時（ja-JP形式）　変更者名  /  summaryテキスト（複数行対応）',
    '件数はセクション見出しにバッジ表示（例：「変更履歴（直近） [3件]」）',
    'リスト最大高さ：288px — 超えた場合はスクロール',
    '履歴ゼロ時：「履歴はまだありません。」を表示',
    '保持上限：最大30件（keepLatestCount: 30）、超過分は古いものから自動削除',
]
for item in layout_items:
    p = doc.add_paragraph(item, style='List Bullet')
    p.runs[0].font.size = Pt(10)

doc.add_paragraph()

# Example
add_heading(doc, '表示イメージ', level=3)

example_text = (
    "変更履歴（直近）  [3件]\n"
    "────────────────────────────────────────\n"
    "⊕  2025/04/01 10:00  田中 太郎\n"
    "   新規作成: タイトル=\"春の新商品\" / 商品件数=1\n"
    "────────────────────────────────────────\n"
    "✎  2025/04/02 14:30  田中 太郎\n"
    "   タイトル: 春の新商品 -> 春の新商品キャンペーン\n"
    "   展開スタート月: 4 -> 5\n"
    "────────────────────────────────────────\n"
    "⚠  2025/04/03 09:15  管理者\n"
    "   修正依頼: 商品画像を添付してください"
)
p = doc.add_paragraph()
run = p.add_run(example_text)
run.font.name = 'Courier New'
run.font.size = Pt(9)

out_path = '/home/user/medicom-entry-sheet/docs/エントリーシート詳細画面_仕様書.docx'
doc.save(out_path)
print(f'Saved: {out_path}')
